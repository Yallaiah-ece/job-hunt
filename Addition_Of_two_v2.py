import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import json
import subprocess
import os
import shutil
from datetime import datetime
import threading
import time
import re

class ResumeGeneratorApp:
    def __init__(self, root, default_output_dir="D:/Resumes_Data_Engineers/New_Resumes", default_filename="Yallaiah_Senior_Data_Engineer"):
        self.root = root
        self.root.title("Enhanced Resume Generator")
        self.root.geometry("950x850")  # Slightly increased height
        self.root.minsize(950, 750)  # Set minimum size to ensure buttons are always visible
        self.root.configure(bg='#2c3e50')
        
        # Available font styles
        self.font_styles = [
            'Calibri', 'Arial', 'Times New Roman', 'Georgia', 
            'Verdana', 'Tahoma', 'Trebuchet MS', 'Comic Sans MS'
        ]
        
        # Default values passed as parameters
        self.default_output_dir = default_output_dir
        self.default_filename = default_filename
        self.default_bold_skills_path = "D:/Resumes_Data_Engineers/New_Resumes/bold_keywords.json"

        # Configure modern styles
        self.configure_styles()
        self.create_widgets()
        
    def configure_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Configure color scheme - Modern dark theme
        self.style.configure('TFrame', 
                            background='#34495e',
                            relief='flat')
        
        self.style.configure('TLabel', 
                            background='#34495e',
                            foreground='#ecf0f1',
                            font=('Segoe UI', 10))
        
        self.style.configure('TLabelFrame', 
                            background='#34495e',
                            foreground='#3498db',
                            font=('Segoe UI', 11, 'bold'))
        
        self.style.configure('TLabelFrame.Label',
                            background='#34495e',
                            foreground='#3498db')
        
        self.style.configure('TButton', 
                            background='#3498db',
                            foreground='white',
                            font=('Segoe UI', 10, 'bold'),
                            focuscolor='none',
                            borderwidth=0,
                            relief='flat')
        
        self.style.map('TButton',
                      background=[('active', '#2980b9'),
                                  ('pressed', '#21618c')])
        
        self.style.configure('TEntry', 
                            font=('Segoe UI', 10),
                            fieldbackground='#ecf0f1',
                            borderwidth=1,
                            relief='solid')
        
        self.style.configure('TCombobox',
                            font=('Segoe UI', 10),
                            fieldbackground='#ecf0f1',
                            borderwidth=1,
                            relief='solid')
        
        # Success button style
        self.style.configure('Success.TButton',
                            background='#27ae60',
                            foreground='white',
                            font=('Segoe UI', 10, 'bold'))
        
        self.style.map('Success.TButton',
                      background=[('active', '#229954'),
                                  ('pressed', '#1e8449')])
        
        # Danger button style
        self.style.configure('Danger.TButton',
                            background='#e74c3c',
                            foreground='white',
                            font=('Segoe UI', 10, 'bold'))
        
        self.style.map('Danger.TButton',
                      background=[('active', '#c0392b'),
                                  ('pressed', '#a93226')])
        
        # Processing button style (Orange/Yellow)
        self.style.configure('Processing.TButton',
                            background='#f39c12',
                            foreground='white',
                            font=('Segoe UI', 10, 'bold'))
        
        self.style.map('Processing.TButton',
                      background=[('active', '#e67e22'),
                                  ('pressed', '#d35400')])
        
    def create_widgets(self):
        # Main frame with modern styling
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title label
        title_label = ttk.Label(main_frame, 
                               text="Professional Resume Generator", 
                               font=('Segoe UI', 16, 'bold'),
                               foreground='#3498db')
        title_label.pack(pady=(0, 20))
        
        # Configuration section
        config_frame = ttk.LabelFrame(main_frame, text="📝 Configuration", padding="15")
        config_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Output directory section
        dir_frame = ttk.Frame(config_frame)
        dir_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(dir_frame, text="Output Directory:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.output_dir_var = tk.StringVar(value=self.default_output_dir) # Use default_output_dir
        self.output_dir_entry = ttk.Entry(dir_frame, textvariable=self.output_dir_var, width=50)
        self.output_dir_entry.grid(row=0, column=1, padx=(0, 10), sticky=tk.W)
        
        ttk.Button(dir_frame, text="📁 Browse", command=self.browse_output_directory).grid(row=0, column=2)
        
        # File name section
        file_frame = ttk.Frame(config_frame)
        file_frame.pack(fill=tk.X, pady=(5, 5))
        
        ttk.Label(file_frame, text="File Name:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.filename_var = tk.StringVar(value=self.default_filename) # Use default_filename
        self.filename_entry = ttk.Entry(file_frame, textvariable=self.filename_var, width=40)
        self.filename_entry.grid(row=0, column=1, padx=(0, 10), sticky=tk.W)
        
        # Format selection
        ttk.Label(file_frame, text="Format:").grid(row=0, column=2, sticky=tk.W, padx=(20, 10))
        self.format_var = tk.StringVar(value="Both (DOCX + PDF)")
        self.format_combo = ttk.Combobox(file_frame, textvariable=self.format_var, 
                                         values=["DOCX Only", "PDF Only", "Both (DOCX + PDF)"], 
                                         state="readonly", width=18)
        self.format_combo.grid(row=0, column=3, padx=(0, 10), sticky=tk.W)
        
        # Font selection section
        font_frame = ttk.Frame(config_frame)
        font_frame.pack(fill=tk.X, pady=(10, 5))
        
        ttk.Label(font_frame, text="Font Style:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.font_var = tk.StringVar(value="Calibri")
        self.font_combo = ttk.Combobox(font_frame, textvariable=self.font_var, 
                                         values=self.font_styles, state="readonly", width=20)
        self.font_combo.grid(row=0, column=1, padx=(0, 10), sticky=tk.W)
        
        ttk.Label(font_frame, text="Font Size:").grid(row=0, column=2, sticky=tk.W, padx=(20, 10))
        self.font_size_var = tk.StringVar(value="11")
        font_size_combo = ttk.Combobox(font_frame, textvariable=self.font_size_var,
                                         values=['9', '10', '11', '12', '13', '14'], 
                                         state="readonly", width=8)
        font_size_combo.grid(row=0, column=3, sticky=tk.W)
        
        # Bold Skills JSON Path section
        bold_skills_frame = ttk.Frame(config_frame)
        bold_skills_frame.pack(fill=tk.X, pady=(5, 0))
        
        ttk.Label(bold_skills_frame, text="Bold Skills JSON:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.bold_skills_var = tk.StringVar(value=self.default_bold_skills_path)  # Set default path
        self.bold_skills_entry = ttk.Entry(bold_skills_frame, textvariable=self.bold_skills_var, width=50)
        self.bold_skills_entry.grid(row=0, column=1, padx=(0, 10), sticky=tk.W)
        
        ttk.Button(bold_skills_frame, text="📁 Browse", command=self.browse_bold_skills_file).grid(row=0, column=2)
        
        # JSON input section with modern styling
        json_frame = ttk.LabelFrame(main_frame, text="📄 Resume JSON Data", padding="15")
        json_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Text widget with custom styling
        text_frame = tk.Frame(json_frame, bg='#34495e')
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        self.json_text = tk.Text(text_frame, 
                                 width=80, height=15,  # Reduced height from 18 to 15
                                 wrap=tk.WORD, 
                                 font=('Consolas', 10),
                                 bg='#2c3e50',
                                 fg='#ecf0f1',
                                 insertbackground='#3498db',
                                 selectbackground='#3498db',
                                 selectforeground='white',
                                 relief='flat',
                                 borderwidth=2)
        
        # Scrollbar for text widget
        scrollbar = tk.Scrollbar(text_frame, bg='#34495e', troughcolor='#2c3e50')
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.json_text.pack(fill=tk.BOTH, expand=True)
        self.json_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.json_text.yview)
        
        # Buttons frame with improved layout - Fixed positioning
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=(5, 5))  # Reduced top padding and ensured it's packed properly
        
        # Left side buttons
        left_buttons = ttk.Frame(buttons_frame)
        left_buttons.pack(side=tk.LEFT)
        
        ttk.Button(left_buttons, text="📂 Load JSON", command=self.load_json_file).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(left_buttons, text="🗑️ Clear", command=self.clear_fields, 
                           style='Danger.TButton').pack(side=tk.LEFT, padx=(0, 10))
        
        # Right side buttons
        right_buttons = ttk.Frame(buttons_frame)
        right_buttons.pack(side=tk.RIGHT)
        
        ttk.Button(right_buttons, text="❌ Exit", command=self.root.quit,
                           style='Danger.TButton').pack(side=tk.RIGHT, padx=(10, 0))
        
        # Store reference to generate button for style changes
        self.generate_button = ttk.Button(right_buttons, text="🚀 Generate Resume", 
                                        command=self.generate_resume_threaded,
                                        style='Success.TButton')
        self.generate_button.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Status bar with progress - Fixed positioning
        status_frame = ttk.Frame(main_frame)
        status_frame.pack(fill=tk.X, pady=(5, 10))  # Ensured proper spacing
        
        self.status_var = tk.StringVar(value="Ready to generate resume...")
        self.status_bar = ttk.Label(status_frame, textvariable=self.status_var, 
                                     font=('Segoe UI', 9), foreground='#7f8c8d')
        self.status_bar.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Progress bar
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate', length=200)
        self.progress.pack(side=tk.RIGHT, padx=(10, 0))

        # --- Contact Information Label ---
        contact_info_label = ttk.Label(main_frame,
                                       text="For any help, contact Yallaiah Onteru || Onteruyallaiah970@gmail.com. @all the rights are reserved.",
                                       font=('Segoe UI', 8),
                                       foreground='#7f8c8d',
                                       anchor=tk.E)
        contact_info_label.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 0))
        
    def browse_output_directory(self):
        initial_dir = self.output_dir_var.get()
        if not os.path.exists(initial_dir):
            initial_dir = os.path.expanduser("~/Documents")
            
        directory = filedialog.askdirectory(
            initialdir=initial_dir,
            title="Select Output Directory"
        )
        if directory:
            self.output_dir_var.set(directory)
            self.status_var.set(f"Output directory set: {directory}")
    
    def browse_bold_skills_file(self):
        """Browse for bold skills JSON file"""
        initial_dir = os.path.dirname(self.bold_skills_var.get()) if self.bold_skills_var.get() else self.default_output_dir
        
        file_path = filedialog.askopenfilename(
            initialdir=initial_dir,
            title="Select Bold Skills JSON File",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if file_path:
            self.bold_skills_var.set(file_path)
            self.status_var.set(f"Bold skills file selected: {os.path.basename(file_path)}")
    
    def load_bold_skills(self, file_path):
        """Load bold skills from JSON file"""
        try:
            if not file_path or not os.path.exists(file_path):
                return []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if 'skills' in data and isinstance(data['skills'], list):
                    return data['skills']
                else:
                    messagebox.showwarning("Warning", "Bold skills JSON file should contain a 'skills' array")
                    return []
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load bold skills file: {str(e)}")
            return []
    
    def load_json_file(self):
        file_path = filedialog.askopenfilename(
            title="Select JSON Resume File",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    json_data = f.read()
                    self.json_text.delete(1.0, tk.END)
                    self.json_text.insert(tk.END, json_data)
                    self.status_var.set(f"Loaded: {os.path.basename(file_path)}")
                    
                    # Update filename based on title in JSON
                    self.update_filename_from_json(json_data)
                    
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load JSON file: {str(e)}")
                self.status_var.set("Error loading file")
    
    def update_filename_from_json(self, json_data):
        """Update filename based on title from JSON data"""
        try:
            data = json.loads(json_data)
            title = data.get('title', '').strip()
            
            if title:
                # Clean title for filename (remove special characters)
                clean_title = re.sub(r'[^\w\s-]', '', title)
                clean_title = re.sub(r'[\s]+', '_', clean_title)
                new_filename = f"Yallaiah_{clean_title}"
            else:
                new_filename = "Yallaiah_Senior_Data_Engineer"
            
            self.filename_var.set(new_filename)
            self.status_var.set(f"Filename updated to: {new_filename}")
            
        except json.JSONDecodeError:
            # If JSON is invalid, keep current filename
            pass
        except Exception as e:
            print(f"Error updating filename: {e}")
    
    def clear_fields(self):
        self.json_text.delete(1.0, tk.END)
        self.filename_var.set("Yallaiah_Senior_Data_Engineer")
        self.font_var.set("Calibri")
        self.font_size_var.set("11")
        self.format_var.set("Both (DOCX + PDF)")
        self.bold_skills_var.set(self.default_bold_skills_path)  # Reset to default path
        self.status_var.set("Fields cleared")
    
    def convert_docx_to_pdf_multiple_methods(self, docx_path):
        """Try multiple methods to convert DOCX to PDF"""
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        
        # Method 1: Try MS Word COM automation (if available)
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
            doc.Close()
            word.Quit()
            print(f"✅ PDF created using MS Word COM: {pdf_path}")
            return True, pdf_path
        except Exception as e:
            print(f"❌ MS Word COM failed: {e}")
            try:
                word.Quit()
            except:
                pass
        
        # Method 2: Try LibreOffice command line
        try:
            output_dir = os.path.dirname(docx_path)
            cmd = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(pdf_path):
                print(f"✅ PDF created using LibreOffice: {pdf_path}")
                return True, pdf_path
            else:
                print(f"❌ LibreOffice conversion failed: {result.stderr}")
        except subprocess.TimeoutExpired:
            print("❌ LibreOffice conversion timed out")
        except FileNotFoundError:
            print("❌ LibreOffice not found in PATH")
        except Exception as e:
            print(f"❌ LibreOffice conversion error: {e}")
        
        # Method 3: Try alternative LibreOffice paths
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        ]
        
        for path in libreoffice_paths:
            if os.path.exists(path):
                try:
                    output_dir = os.path.dirname(docx_path)
                    cmd = [
                        path, "--headless", "--convert-to", "pdf",
                        "--outdir", output_dir, docx_path
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0 and os.path.exists(pdf_path):
                        print(f"✅ PDF created using LibreOffice at {path}: {pdf_path}")
                        return True, pdf_path
                    else:
                        print(f"❌ LibreOffice at {path} failed: {result.stderr}")
                except Exception as e:
                    print(f"❌ LibreOffice at {path} error: {e}")
        
        # Method 4: Try python-docx2pdf (if available)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                print(f"✅ PDF created using docx2pdf: {pdf_path}")
                return True, pdf_path
        except ImportError:
            print("❌ docx2pdf not available")
        except Exception as e:
            print(f"❌ docx2pdf conversion failed: {e}")
        
        print("❌ All PDF conversion methods failed")
        return False, None
    
    def make_text_bold_for_skills(self, paragraph, text, bold_skills, font_name, font_size):
        """Add text to paragraph with bold formatting for matching skills"""
        if not bold_skills or not text:
            # If no bold skills or empty text, add normal text
            run = paragraph.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            return
        
        # Create a case-insensitive pattern for whole word matching
        # Sort skills by length (longest first) to avoid partial matches
        sorted_skills = sorted(bold_skills, key=len, reverse=True)
        
        # Create pattern that matches whole words only
        skill_patterns = []
        for skill in sorted_skills:
            # Escape special regex characters and create word boundary pattern
            escaped_skill = re.escape(skill)
            skill_patterns.append(f"\\b{escaped_skill}\\b")
        
        if not skill_patterns:
            run = paragraph.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            return
        
        # Combine all patterns with OR
        pattern = "|".join(skill_patterns)
        
        # Find all matches
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        
        if not matches:
            # No matches found, add normal text
            run = paragraph.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            return
        
        # Process text with matches
        last_end = 0
        
        for match in matches:
            start, end = match.span()
            
            # Add text before match (normal)
            if start > last_end:
                before_text = text[last_end:start]
                if before_text:
                    run = paragraph.add_run(before_text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            
            # Add matched text (bold)
            matched_text = text[start:end]
            if matched_text:
                run = paragraph.add_run(matched_text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = True
            
            last_end = end
        
        # Add remaining text after last match (normal)
        if last_end < len(text):
            remaining_text = text[last_end:]
            if remaining_text:
                run = paragraph.add_run(remaining_text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
    
    def generate_resume_threaded(self):
        """Run resume generation in a separate thread to prevent UI freezing"""
        # Change button style to processing (orange)
        self.generate_button.configure(style='Processing.TButton')
        self.generate_button.configure(text="⏳ Processing...")
        
        self.progress.start(10)
        self.status_var.set("Generating resume...")
        
        thread = threading.Thread(target=self.generate_resume)
        thread.daemon = True
        thread.start()
    
    def reset_generate_button(self):
        """Reset generate button to original state"""
        self.generate_button.configure(style='Success.TButton')
        self.generate_button.configure(text="🚀 Generate Resume")
        self.progress.stop()
    
    def generate_resume(self):
        json_string = self.json_text.get(1.0, tk.END).strip()
        filename = self.filename_var.get().strip()
        output_dir = self.output_dir_var.get().strip()
        selected_format = self.format_var.get()
        selected_font = self.font_var.get()
        font_size = int(self.font_size_var.get())
        bold_skills_path = self.bold_skills_var.get().strip()
        
        if not json_string:
            self.reset_generate_button()
            messagebox.showerror("Error", "Please provide JSON resume data")
            self.status_var.set("Error: No JSON data provided")
            return
        
        if not filename:
            self.reset_generate_button()
            messagebox.showerror("Error", "Please specify a filename")
            self.status_var.set("Error: No filename specified")
            return
        
        if not output_dir:
            self.reset_generate_button()
            messagebox.showerror("Error", "Please specify an output directory")
            self.status_var.set("Error: No output directory specified")
            return
        
        try:
            # Validate JSON and update filename
            json_data = json.loads(json_string)
            
            # Update filename based on title if not already done
            title = json_data.get('title', '').strip()
            if title and not filename.startswith(f"Yallaiah_{title.replace(' ', '_')}"):
                clean_title = re.sub(r'[^\w\s-]', '', title)
                clean_title = re.sub(r'[\s]+', '_', clean_title)
                filename = f"Yallaiah_{clean_title}"
                self.filename_var.set(filename)
            elif not title and filename == self.default_filename:
                filename = "Yallaiah_Senior_Data_Engineer"
                self.filename_var.set(filename)
            
            # Load bold skills
            bold_skills = self.load_bold_skills(bold_skills_path) if bold_skills_path else []
            
            # Create output directory if it doesn't exist
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate resume with selected format and bold skills
            success = self.generate_resume_from_json(json_string, filename, output_dir, selected_format, selected_font, font_size, bold_skills)
            
            self.reset_generate_button()
            
            if success:
                messagebox.showinfo("Success", f"Resume generated successfully!\nFiles saved to: {output_dir}")
                self.status_var.set("✅ Resume generated successfully!")
            else:
                messagebox.showerror("Error", "Failed to generate resume")
                self.status_var.set("❌ Failed to generate resume")
                
        except json.JSONDecodeError:
            self.reset_generate_button()
            messagebox.showerror("Error", "Invalid JSON format")
            self.status_var.set("Error: Invalid JSON format")
        except Exception as e:
            self.reset_generate_button()
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            self.status_var.set(f"Error: {str(e)}")
    
    def generate_resume_from_json(self, json_string, filename, output_dir, selected_format, font_name, font_size, bold_skills):
        try:
            data = json.loads(json_string)
            
            # Use the exact output directory specified (no date-based subdirectory)
            os.makedirs(output_dir, exist_ok=True)
            
            # Determine file paths
            docx_filename = f"{filename}.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            
            # Generate DOCX if needed
            if selected_format in ["DOCX Only", "Both (DOCX + PDF)"]:
                doc = Document()

                # === Styling & Layout with dynamic font ===
                section = doc.sections[0]
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)

                # Apply selected font style
                style = doc.styles['Normal']
                font = style.font
                font.name = font_name
                font.size = Pt(font_size)
                style.paragraph_format.line_spacing = 1.0
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                def add_centered_paragraph(text, bold=False, size=None):
                    if size is None:
                        size = font_size
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(text)
                    run.bold = bold
                    run.font.name = font_name
                    run.font.size = Pt(size)
                    p.paragraph_format.space_after = Pt(2)

                def add_section_heading(text):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(text.upper())
                    run.bold = True
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p_border = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), '000000')
                    p_border.append(bottom)
                    p._p.get_or_add_pPr().append(p_border)
                    p.paragraph_format.space_after = Pt(4)

                def add_bullet_points(items):
                    """Add bullet points with bold skills formatting"""
                    for item in items:
                        p = doc.add_paragraph(style='List Bullet')
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.left_indent = Inches(0.25)
                        
                        # Use the new function to add text with bold skills
                        self.make_text_bold_for_skills(p, item, bold_skills, font_name, font_size)

                def add_hyperlinked_paragraph(doc, text_parts):
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(4)
                    for idx, (display_text, url) in enumerate(text_parts):
                        r_id = doc.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                        hyperlink = OxmlElement('w:hyperlink')
                        hyperlink.set(qn('r:id'), r_id)
                        new_run = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        
                        # Apply font to hyperlink
                        font_elem = OxmlElement('w:rFonts')
                        font_elem.set(qn('w:ascii'), font_name)
                        font_elem.set(qn('w:hAnsi'), font_name)
                        rPr.append(font_elem)
                        
                        color = OxmlElement('w:color')
                        color.set(qn('w:val'), '0000FF')
                        rPr.append(color)
                        underline = OxmlElement('w:u')
                        underline.set(qn('w:val'), 'single')
                        rPr.append(underline)
                        new_run.append(rPr)
                        text = OxmlElement('w:t')
                        text.text = display_text
                        new_run.append(text)
                        hyperlink.append(new_run)
                        paragraph._p.append(hyperlink)
                        if idx != len(text_parts) - 1:
                            run = paragraph.add_run(" | ")
                            run.font.name = font_name

                # === HEADER ===
                add_centered_paragraph(data['name'], bold=True, size=font_size + 3)
                add_centered_paragraph(data.get('title', ''), size=font_size)

                contact = data.get('contact', {})
                contact_parts = []
                if contact.get('portfolio'):
                    contact_parts.append(('Portfolio', contact['portfolio']))
                if contact.get('linkedin'):
                    contact_parts.append(('LinkedIn', contact['linkedin']))
                if contact.get('email'):
                    contact_parts.append((contact['email'], f"mailto:{contact['email']}"))
                if contact.get('phone'):
                    contact_parts.append((contact['phone'], f"tel:{contact['phone']}"))

                # Handle legacy format
                if data.get('portfolio'):
                    contact_parts.append(('Portfolio', data['portfolio']))
                if data.get('linkedin'):
                    contact_parts.append(('LinkedIn', data['linkedin']))
                if data.get('email') and not any(part[0] == data['email'] for part in contact_parts):
                    contact_parts.append((data['email'], f"mailto:{data['email']}"))
                if data.get('phone') and not any(part[0] == data['phone'] for part in contact_parts):
                    contact_parts.append((data['phone'], f"tel:{data['phone']}"))

                if contact_parts:
                    add_hyperlinked_paragraph(doc, contact_parts)

                # === PROFESSIONAL SUMMARY ===
                if data.get('professional_summary'):
                    add_section_heading("Professional Summary")
                    add_bullet_points(data['professional_summary'])  # This will use bold skills

                # === TECHNICAL SKILLS ===
                if data.get('technical_skills'):
                    add_section_heading("Technical Skills")
                    for category, skills in data['technical_skills'].items():
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(f"• {category}: ")
                        run.bold = True
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        skill_run = p.add_run(", ".join(skills))
                        skill_run.font.name = font_name
                        skill_run.font.size = Pt(font_size)

                # === EXPERIENCE ===
                if data.get('experience'):
                    add_section_heading("Professional Experience")
                    for job in data['experience']:
                        p = doc.add_paragraph()
                        run = p.add_run(f"Role: {job['role']}")
                        run.bold = True
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        p.paragraph_format.space_after = Pt(0)

                        p = doc.add_paragraph()
                        p.paragraph_format.tab_stops.clear_all()
                        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3))
                        run_left = p.add_run(f"Client: {job['company']}")
                        run_left.bold = True
                        run_left.font.name = font_name
                        run_left.font.size = Pt(font_size)
                        if job.get('duration'):
                            run_right = p.add_run(f"\t{job['duration']}")
                            run_right.bold = True
                            run_right.font.name = font_name
                            run_right.font.size = Pt(font_size - 1)
                            run_right.font.color.rgb = RGBColor(0, 0, 0)
                        p.paragraph_format.space_after = Pt(4)

                        if job.get('project_overview'):
                            p = doc.add_paragraph()
                            run = p.add_run("Project Overview: ")
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            desc_run = p.add_run(job['project_overview'])
                            desc_run.font.name = font_name
                            desc_run.font.size = Pt(font_size)
                            p.paragraph_format.space_after = Pt(4)

                        if job.get('responsibilities'):
                            p = doc.add_paragraph()
                            run = p.add_run("Responsibilities: ")
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            p.paragraph_format.space_after = Pt(2)
                            add_bullet_points(job['responsibilities'])  # This will use bold skills

                        if job.get('environment'):
                            p = doc.add_paragraph()
                            run = p.add_run("Environment: ")
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            env_run = p.add_run(", ".join(job['environment']))
                            env_run.font.name = font_name
                            env_run.font.size = Pt(font_size)
                            p.paragraph_format.space_after = Pt(8)

                # === EDUCATION ===
                if data.get('education') and isinstance(data['education'], dict):
                    add_section_heading("Education")
                    edu = data['education']
                    p = doc.add_paragraph()
                    line_parts = []
                    if edu.get('degree'):
                        line_parts.append(edu['degree'])
                    if edu.get('field'):
                        line_parts.append(edu['field'])
                    if edu.get('institution'):
                        line_parts.append(f"at {edu['institution']}")
                    if edu.get('year'):
                        line_parts.append(f"({edu['year']})")
                    if line_parts:
                        run = p.add_run(", ".join(line_parts))
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                    p.paragraph_format.space_after = Pt(2)

                # === CERTIFICATIONS ===
                if data.get('certifications'):
                    add_section_heading("Certifications")
                    for cert in data['certifications']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(cert)
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        p.paragraph_format.space_after = Pt(2)

                # Save DOCX
                doc.save(docx_path)
                print(f"✅ DOCX saved to: {docx_path}")

            # Convert to PDF if needed
            if selected_format in ["PDF Only", "Both (DOCX + PDF)"]:
                if selected_format == "PDF Only" and not os.path.exists(docx_path):
                    # Generate DOCX temporarily for PDF conversion
                    self.generate_resume_from_json(json_string, filename, output_dir, "DOCX Only", font_name, font_size, bold_skills)
                
                # Convert using multiple methods
                pdf_success, pdf_path = self.convert_docx_to_pdf_multiple_methods(docx_path)
                
                if not pdf_success:
                    print("⚠️ PDF conversion failed")
                    if selected_format == "PDF Only":
                        return False
                else:
                    print(f"✅ PDF saved to: {pdf_path}")
                
                # Remove temporary DOCX if PDF Only was selected
                if selected_format == "PDF Only" and os.path.exists(docx_path):
                    os.remove(docx_path)
                    print("✅ Temporary DOCX file removed")

            print(f"✅ Files saved to: {output_dir}")
            return True
            
        except Exception as e:
            print(f"❌ Error generating resume: {e}")
            return False

if __name__ == "__main__":
    root = tk.Tk()
    
    # Define your dynamic parameters here
    DEFAULT_OUTPUT_DIRECTORY = "D:/Resumes_Data_Engineers/New_Resumes"
    DEFAULT_FILE_NAME = "Yallaiah_Senior_Data_Engineer"  # Updated default filename

    app = ResumeGeneratorApp(root, default_output_dir=DEFAULT_OUTPUT_DIRECTORY, default_filename=DEFAULT_FILE_NAME)
    root.mainloop()
